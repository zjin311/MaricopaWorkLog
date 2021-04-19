#%%
'''test python-docx func'''
import docx
path_doc='C:/Users/Mr.Goldss/Desktop/switch back to risk management/return to work program'
doc = docx.Document(path_doc+\
    '/Only Loss Days analysis draft.docx')


# read all paragraphs
all_paras = doc.paragraphs
len(all_paras)

for para in all_paras:
    print(para.text)

# read text in a specific paragraph
single_para = doc.paragraphs[5]
print(single_para.text)

#run 的定义是一句话里边不同的字体，通常是重点 加粗 或者加下划线的
single_para = doc.paragraphs[5]
for run in single_para.runs:
    print(run.text)


#%%
#This section is about how to write a docx file, which is what I need
#step 1: create a file
mydoc = docx.Document()
#step 2: add_paragraph
mydoc.add_paragraph("This is first paragraph of a MS Word file.")
#step 3: save your file
mydoc.save('C:/Users/Mr.Goldss/Desktop/my_written_file.docx')
#%%
#use the following code to add another paragraph and then save to the same address
mydoc.add_paragraph("i am trying to add another para@34234234234")
#step 3: save your file (save address + file name)
mydoc.save('C:/Users/Mr.Goldss/Desktop/my_written_file.docx')

#%%
third_para = mydoc.add_paragraph("This is the 3th paragraph.")
third_para.add_run(" this is a section at the end of third paragraph")
mydoc.save('C:/Users/Mr.Goldss/Desktop/my_written_file.docx')

#%%
#adding header
mydoc.add_heading("This is level 1 heading", 0)
mydoc.add_heading("This is level 2 heading", 1)
mydoc.add_heading("This is level 3 heading", 2)
mydoc.save('C:/Users/Mr.Goldss/Desktop/my_written_file.docx')
#%%
# adding image
mydoc.add_picture('C:/Users/Mr.Goldss/Pictures/selfie gold wedding.jpg')
mydoc.save('C:/Users/Mr.Goldss/Desktop/my_written_file.docx')
